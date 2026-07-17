from __future__ import annotations

from docx import Document

from src import docx_preprocessor, main


def test_docx_chunk_records_preserve_heading_sections(tmp_path) -> None:
    path = tmp_path / "report.docx"
    document = Document()
    document.add_heading("Executive summary", level=1)
    document.add_paragraph("Revenue increased while share remained flat.")
    document.add_heading("Safety", level=1)
    document.add_paragraph("No new safety signal was identified.")
    document.save(path)

    records = docx_preprocessor.extract_docx_chunk_records(path)

    assert [record.section_title for record in records] == [
        "Executive summary",
        "Safety",
    ]
    assert "Revenue increased" in records[0].text
    assert "No new safety signal" in records[1].text


def test_context_projects_pdf_page_docx_section_and_pptx_slide_locations() -> None:
    context, sources, empty_pages = main._context_from_hits(
        [
            {
                "text": "The primary endpoint was met.",
                "summary": '{"source_channel":"native_text"}',
                "doc_id": 6,
                "file_name": "study.pdf",
                "i_page": 4,
                "i_chunk_on_doc": 1,
                "_additional": {"id": "pdf-4", "distance": 0.05},
            },
            {
                "text": "No new safety signal was identified.",
                "summary": '{"source_channel":"native_text","section_title":"Safety"}',
                "doc_id": 7,
                "file_name": "report.docx",
                "i_page": 1,
                "i_chunk_on_doc": 2,
                "_additional": {"id": "docx-2", "distance": 0.1},
            },
            {
                "text": "The market outlook remains positive.",
                "summary": "{}",
                "doc_id": 8,
                "file_name": "brief.pptx",
                "i_page": 12,
                "i_chunk_on_doc": 11,
                "_additional": {"id": "pptx-12", "distance": 0.2},
            },
        ]
    )

    assert empty_pages == []
    assert "page=4" in context
    assert "section=Safety" in context
    assert "slide=12" in context
    assert sources[0].i_page == 4
    assert sources[0].slide_number is None
    assert sources[1].section_title == "Safety"
    assert sources[1].slide_number is None
    assert sources[2].section_title is None
    assert sources[2].slide_number == 12
